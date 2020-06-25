import docx
""""
d = docx.Document('F:\\python_execsice\\reading word\\DN-DYVN.docx')

print(d.paragraphs[0].text)

d.add_paragraph("Add one lines")
d.save('F:\\python_execsice\\reading word\\DN-DYVN.docx')
"""
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('F:\\python_execsice\\reading word\\DN-DYVN.docx'))
