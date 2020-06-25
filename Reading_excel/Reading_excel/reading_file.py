# Download example.xlsx from :http://autbor.com/example.xlsx
# When we work with excel file.We have some concepts:
#workbook = file excel, sheet means in excel file include many sheets
#row and columm, cell
#


import openpyxl
import os
from docx import Document
import pandas as pd
#os.chdir('D:\\Python_lesson\\python_execsice\\Reading_excel')

#workbook = openpyxl.load_workbook('example.xlsx')

document = Document('DN-DYVN.docx')


#sheet = workbook.get_sheet_by_name('Sheet1')
#sheet = workbook["Input data"]
#sheet['B3'] = "Hello"
#sheet["C3"] = "World"
#workbook.save('example.xlsx')

table = document.tables[1]

part_no = []
part_name = []
quanity_product = []
quanity_box = []
row_index = 0
for row in table.rows:

    print("Length of row.cells",len(row.cells))
    #print("type of row.cells", type(row.cells))
    for cell in row.cells:
        if row_index == 1:
            #for para in cell.paragraphs:
                #if len(para.text) != 0 :
            part_no.append(cell.paragraphs[0].text)
        if row_index == 2:
            #for para in cell.paragraphs:
                #if len(para.text) != 0 :
            part_name.append(cell.paragraphs[0].text)
        if row_index == 3:
            #for para in cell.paragraphs:
                #if len(para.text) != 0 :
            quanity_product.append(cell.paragraphs[0].text)

        if row_index == 4:
            #for para in cell.paragraphs:
                #if len(para.text) != 0 :
            quanity_box.append(cell.paragraphs[0].text)

        row_index = row_index + 1
    row_index = 0


print(part_no)
print(part_name)
print(quanity_product)
print(quanity_box)

""""
workbook = openpyxl.load_workbook('example.xlsx')
sheet = workbook["Sheet1"]
"""
"""
sheet['B3'] = "Hello"
sheet["C3"] = "World"
"""
import xlsxwriter
workbook = xlsxwriter.Workbook('example.xlsx')
#sheet = workbook["Sheet1"]
#sheet =  workbook.worksheets("Sheet1")
workbook.add_worksheet(Sheet1)
success = workbook.add_format(
    {
        "bg_color":"#00ff00",
        "font":"Century",
        "font_size":22
    
    }
)

current_index = 1
for part_no_index in range(1,len(part_no)-1):
    if len(part_no[part_no_index]) == 0:
        current_index -=1
    else:
        sheet['C' + str(current_index + 11)] = part_no[part_no_index]
        sheet['K' + str(current_index + 11)] = part_name[part_no_index]
        sheet['L' + str(current_index + 11)] = quanity_product[part_no_index]
    current_index +=1



    

workbook.save('example.xlsx')

"""
type(sheet)

cell = sheet['A1']
print(cell.value)
print(sheet['A2'].value)

print(sheet['B1'].value)

print('Display the first columm A:')

for i in range(1,8):
    print(sheet['A' + str(i)].value)


print('Display the first columm B:')

for i in range(1,8):
    print(sheet['B' + str(i)].value)
    
#How to work with a cell
print(sheet.cell(row = 1, column = 2).value)
"""
